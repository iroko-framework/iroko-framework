#!/usr/bin/env python3
"""
docx_to_html.py  --  Iroko Historical Society
Converts any .docx file to a styled HTML document matching the Iroko whitepaper aesthetic.

Usage:
    python3 docx_to_html.py input.docx [output.html]
    python3 docx_to_html.py input.docx --title "Custom Title" --subtitle "Subtitle" --version "v1.0"

If output filename is omitted, it is derived from the input filename.

Required:
    pip install python-docx --break-system-packages
"""

import sys
import os
import re
import html as html_mod
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import RGBColor
    from docx.oxml.ns import qn
except ImportError:
    print("ERROR: python-docx is required. Run: pip install python-docx --break-system-packages")
    sys.exit(1)


# ============================================================
# HELPERS
# ============================================================

def slugify(text):
    """Convert a heading to a URL-safe anchor id."""
    text = text.lower().strip()
    text = re.sub(r'[^\w\s-]', '', text)
    text = re.sub(r'[\s_-]+', '-', text)
    text = re.sub(r'^-+|-+$', '', text)
    return text or "section"


def escape(text):
    return html_mod.escape(text or "")


def is_code_block(para):
    """Detect code/pre content: Courier New font or style name contains 'code'."""
    style_name = ((para.style.name if para.style else "") or "").lower()
    if "code" in style_name or "preformat" in style_name:
        return True
    for run in para.runs:
        font_name = (run.font.name or "").lower()
        if "courier" in font_name or "consolas" in font_name or "monospace" in font_name:
            return True
    return False


def run_to_html(run):
    """Convert a single docx run to inline HTML."""
    text = escape(run.text)
    if not text:
        return ""
    # Bold
    if run.bold:
        text = f"<strong>{text}</strong>"
    # Italic
    if run.italic:
        text = f"<em>{text}</em>"
    # Monospace font
    font_name = (run.font.name or "").lower()
    if "courier" in font_name or "consolas" in font_name:
        text = f"<code>{text}</code>"
    return text


def para_inline_html(para):
    """Convert all runs in a paragraph to inline HTML."""
    parts = []
    for run in para.runs:
        parts.append(run_to_html(run))
    return "".join(parts)


def get_heading_level(para):
    """Return 1-6 for heading styles, 0 otherwise."""
    style = (para.style.name if para.style else "") or ""
    m = re.match(r'Heading\s+(\d)', style, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return 0


def is_bullet(para):
    style = ((para.style.name if para.style else "") or "").lower()
    return ("list" in style or "bullet" in style
            or para._element.find(qn('w:numPr')) is not None)


def is_numbered(para):
    style = ((para.style.name if para.style else "") or "").lower()
    return "number" in style or "list number" in style


def is_horizontal_rule(para):
    """Detect a paragraph that is just dashes or underscores (manual divider)."""
    text = para.text.strip()
    return bool(re.match(r'^[-_=]{3,}$', text))


def detect_note(para):
    """Return note type ('note', 'warning', 'tip') or None."""
    text = para.text.strip()
    lower = text.lower()
    if lower.startswith("note:") or lower.startswith("note "):
        return "note"
    if lower.startswith("warning:") or lower.startswith("warn:"):
        return "warning"
    if lower.startswith("tip:") or lower.startswith("tip "):
        return "tip"
    return None


def table_to_html(table):
    """Convert a docx Table to an HTML table."""
    rows = table.rows
    if not rows:
        return ""
    lines = ['<div class="module-table-wrapper"><table class="module-table">']
    # First row as header
    first = rows[0]
    lines.append("<thead><tr>")
    for cell in first.cells:
        text = escape(cell.text.strip())
        lines.append(f"<th>{text}</th>")
    lines.append("</tr></thead><tbody>")
    for row in rows[1:]:
        lines.append("<tr>")
        for cell in row.cells:
            text = escape(cell.text.strip())
            lines.append(f"<td>{text}</td>")
        lines.append("</tr>")
    lines.append("</tbody></table></div>")
    return "\n".join(lines)


# ============================================================
# MAIN PARSER
# ============================================================

class DocxConverter:
    def __init__(self, docx_path, title=None, subtitle=None, version=None,
                 author=None, org=None, doc_type=None):
        self.doc = Document(docx_path)
        self.basename = Path(docx_path).stem

        # Collect all headings for TOC before rendering body
        self.headings = []          # list of (level, text, slug, counter_label)
        self.h1_count = 0
        self.h2_count = 0
        self.h3_count = 0

        # Meta overrides
        self.title    = title    or self._guess_title()
        self.subtitle = subtitle or ""
        self.version  = version  or ""
        self.author   = author   or "Iroko Historical Society"
        self.org      = org      or "Iroko Historical Society"
        self.doc_type = doc_type or "Technical Manual"

        # Pre-scan headings
        self._scan_headings()

    def _guess_title(self):
        """Use the first Heading 1 or the filename."""
        for para in self.doc.paragraphs:
            if get_heading_level(para) == 1:
                return para.text.strip()
        return self.basename.replace("_", " ").replace("-", " ").title()

    def _scan_headings(self):
        """Pre-scan all headings to build the TOC and assign roman/alpha labels."""
        roman = ["I","II","III","IV","V","VI","VII","VIII","IX","X",
                 "XI","XII","XIII","XIV","XV","XVI","XVII","XVIII","XIX","XX"]
        alpha = list("ABCDEFGHIJKLMNOPQRSTUVWXYZ")
        h1_idx = 0
        h2_idx = 0
        for para in self.doc.paragraphs:
            level = get_heading_level(para)
            text  = para.text.strip()
            if not text:
                continue
            if level == 1:
                label = roman[h1_idx] if h1_idx < len(roman) else str(h1_idx + 1)
                h1_idx += 1
                h2_idx = 0
            elif level == 2:
                label = alpha[h2_idx] if h2_idx < len(alpha) else str(h2_idx + 1)
                h2_idx += 1
            elif level == 3:
                label = ""
            else:
                continue
            slug_base = slugify(text)
            # Deduplicate slugs
            slug = slug_base
            existing = [h[2] for h in self.headings]
            counter = 1
            while slug in existing:
                slug = f"{slug_base}-{counter}"
                counter += 1
            self.headings.append((level, text, slug, label))

    # ----------------------------------------------------------
    # BODY RENDERING
    # ----------------------------------------------------------

    def _render_paragraphs(self):
        """Walk all block-level elements and render to HTML."""
        paras = list(self.doc.paragraphs)
        tables = {id(t): t for t in self.doc.tables}

        # Build a map of paragraph objects to their position in XML,
        # so we can inject tables at the right place.
        # We iterate the document body children directly.
        body_children = list(self.doc.element.body)
        heading_slug_map = {}
        for level, text, slug, label in self.headings:
            heading_slug_map[text] = (slug, label, level)

        html_parts = []
        in_bullets  = False
        in_numbered = False
        in_code     = False
        code_lines  = []

        def flush_list():
            nonlocal in_bullets, in_numbered
            if in_bullets:
                html_parts.append("</ul>")
                in_bullets = False
            if in_numbered:
                html_parts.append("</ol>")
                in_numbered = False

        def flush_code():
            nonlocal in_code, code_lines
            if in_code:
                joined = "\n".join(code_lines)
                html_parts.append(
                    f'<div class="code-block"><pre>{joined}</pre></div>'
                )
                in_code = False
                code_lines = []

        from docx.oxml.ns import qn as _qn
        para_objs = list(self.doc.paragraphs)
        table_objs = list(self.doc.tables)
        para_idx = 0
        table_idx = 0

        for child in body_children:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "tbl":
                # Render table
                flush_list()
                flush_code()
                if table_idx < len(table_objs):
                    html_parts.append(table_to_html(table_objs[table_idx]))
                    table_idx += 1

            elif tag == "p":
                if para_idx >= len(para_objs):
                    para_idx += 1
                    continue
                para = para_objs[para_idx]
                para_idx += 1

                text    = para.text.strip()
                level   = get_heading_level(para)
                inline  = para_inline_html(para)

                # Skip empty paragraphs (use as spacers in lists/code only)
                if not text:
                    if in_code:
                        code_lines.append("")
                    continue

                # Horizontal rule
                if is_horizontal_rule(para):
                    flush_list()
                    flush_code()
                    html_parts.append('<hr class="doc-rule">')
                    continue

                # Headings
                if level in (1, 2, 3):
                    flush_list()
                    flush_code()
                    slug, label, _ = heading_slug_map.get(text, (slugify(text), "", level))
                    if level == 1:
                        html_parts.append(
                            f'<section class="doc-section" id="{slug}">'
                            f'<div class="section-eyebrow">Section {label}</div>'
                            f'<h2 class="section-title">{escape(text)}</h2>'
                        )
                    elif level == 2:
                        html_parts.append(
                            f'<h3 class="subsection-title" id="{slug}">{escape(text)}</h3>'
                        )
                    elif level == 3:
                        html_parts.append(
                            f'<h4 class="subsubsection-title" id="{slug}">{escape(text)}</h4>'
                        )
                    continue

                # Code block (monospace paragraph)
                if is_code_block(para):
                    flush_list()
                    if not in_code:
                        in_code = True
                        code_lines = []
                    code_lines.append(escape(para.text))
                    continue
                else:
                    flush_code()

                # Note / Warning / Tip
                note_type = detect_note(para)
                if note_type:
                    flush_list()
                    cls_map = {"note": "doc-note", "warning": "doc-warning", "tip": "doc-tip"}
                    label_map = {"note": "NOTE", "warning": "WARNING", "tip": "TIP"}
                    cls   = cls_map[note_type]
                    lbl   = label_map[note_type]
                    # Strip the label prefix from the inline content
                    body_html = re.sub(
                        r'^<strong>' + lbl + r'[:\s]*</strong>\s*',
                        '', inline, flags=re.IGNORECASE
                    )
                    body_html = re.sub(
                        r'^' + lbl + r'[:\s]*', '', inline, flags=re.IGNORECASE
                    )
                    html_parts.append(
                        f'<div class="{cls}"><span class="note-label">{lbl}</span> {body_html}</div>'
                    )
                    continue

                # Bullet list
                if is_bullet(para):
                    flush_code()
                    if in_numbered:
                        html_parts.append("</ol>")
                        in_numbered = False
                    if not in_bullets:
                        html_parts.append('<ul class="doc-list">')
                        in_bullets = True
                    html_parts.append(f"<li>{inline}</li>")
                    continue

                # Numbered list
                if is_numbered(para):
                    flush_code()
                    if in_bullets:
                        html_parts.append("</ul>")
                        in_bullets = False
                    if not in_numbered:
                        html_parts.append('<ol class="doc-list doc-list-num">')
                        in_numbered = True
                    html_parts.append(f"<li>{inline}</li>")
                    continue

                # Normal paragraph
                flush_list()
                html_parts.append(f"<p>{inline}</p>")

        flush_list()
        flush_code()

        # Close any open section tags
        # Count open sections by counting h1s vs explicit closes
        open_sections = sum(1 for p in html_parts if 'class="doc-section"' in p)
        html_parts.append("</section>" * open_sections)

        return "\n".join(html_parts)

    # ----------------------------------------------------------
    # TOC
    # ----------------------------------------------------------

    def _render_toc(self):
        items = []
        for level, text, slug, label in self.headings:
            if level > 2:
                continue
            num_html = f'<span class="toc-num">{label}.</span> ' if label else ""
            indent = ' style="padding-left:22px;font-size:12.5px;opacity:0.55;"' if level == 2 else ""
            items.append(
                f'<li><a href="#{slug}"{indent}>{num_html}{escape(text)}</a></li>'
            )
        return "\n".join(items)

    # ----------------------------------------------------------
    # FULL HTML
    # ----------------------------------------------------------

    def render(self):
        body_html = self._render_paragraphs()
        toc_html  = self._render_toc()

        # Version badge
        ver_display = f" &middot; {escape(self.version)}" if self.version else ""
        # Subtitle line
        subtitle_html = (
            f'<div class="cover-subtitle">{escape(self.subtitle)}</div>'
            if self.subtitle else ""
        )

        return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<meta name="author" content="{escape(self.author)}">
<title>{escape(self.title)} — {escape(self.org)}</title>
<link rel="preconnect" href="https://fonts.googleapis.com">
<link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,500;0,600;1,300;1,400;1,500&family=Cinzel:wght@400;500;600&family=Source+Code+Pro:wght@300;400;500&display=swap" rel="stylesheet">
<style>
/* ============================================================
   VARIABLES
   ============================================================ */
:root {{
  --ink:         #1a1208;
  --parchment:   #faf6ee;
  --cream:       #f5f0e6;
  --ochre:       #c8832a;
  --deep-ochre:  #8b5a1a;
  --pale-ochre:  #e8d5a0;
  --sienna:      #7a3d1f;
  --indigo:      #1e2d4a;
  --deep-blue:   #0f1d35;
  --gold:        #d4a843;
  --forest:      #2d4a2d;
  --rule:        rgba(200,131,42,0.25);
  --rule-strong: rgba(200,131,42,0.45);
  --sidebar-w:   300px;
  --max-body:    800px;
}}

/* ============================================================
   RESET & BASE
   ============================================================ */
*, *::before, *::after {{ margin:0; padding:0; box-sizing:border-box; }}
html {{ scroll-behavior: smooth; font-size: 20px; }}
body {{
  background: var(--parchment);
  color: var(--ink);
  font-family: 'Cormorant Garamond', serif;
  font-weight: 400;
  line-height: 1.75;
  min-height: 100vh;
}}
body::before {{
  content: '';
  position: fixed; inset: 0; pointer-events: none; z-index: 0;
  background-image: url("data:image/svg+xml,%3Csvg viewBox='0 0 300 300' xmlns='http://www.w3.org/2000/svg'%3E%3Cfilter id='n'%3E%3CfeTurbulence type='fractalNoise' baseFrequency='0.85' numOctaves='4' stitchTiles='stitch'/%3E%3C/filter%3E%3Crect width='100%25' height='100%25' filter='url(%23n)' opacity='0.035'/%3E%3C/svg%3E");
  opacity: 0.5;
}}

/* ============================================================
   LAYOUT
   ============================================================ */
.layout {{ display: flex; min-height: 100vh; position: relative; z-index: 1; }}

/* ============================================================
   SIDEBAR
   ============================================================ */
.sidebar {{
  width: var(--sidebar-w);
  flex-shrink: 0;
  position: sticky;
  top: 0;
  height: 100vh;
  overflow-y: auto;
  background: #f0ebe0;
  border-right: 1px solid var(--rule);
  padding: 36px 24px 48px;
  display: flex;
  flex-direction: column;
  gap: 0;
}}
.sidebar::-webkit-scrollbar {{ width: 4px; }}
.sidebar::-webkit-scrollbar-track {{ background: transparent; }}
.sidebar::-webkit-scrollbar-thumb {{ background: var(--rule-strong); border-radius: 2px; }}
.sidebar-logo-img {{
  height: 52px; width: auto;
  margin-bottom: 16px;
  opacity: 0.88;
  display: block;
}}
.sidebar-org {{
  font-family: 'Cinzel', serif;
  font-size: 11px;
  font-weight: 400;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  color: var(--ochre);
  margin-bottom: 5px;
}}
.sidebar-title {{
  font-family: 'Cinzel', serif;
  font-size: 14px;
  font-weight: 500;
  letter-spacing: 0.06em;
  color: var(--indigo);
  line-height: 1.4;
  margin-bottom: 4px;
}}
.sidebar-version {{
  font-family: 'Source Code Pro', monospace;
  font-size: 10px;
  letter-spacing: 0.12em;
  color: var(--deep-ochre);
  opacity: 0.6;
  margin-bottom: 28px;
}}
.toc-section-head {{
  font-family: 'Source Code Pro', monospace;
  font-size: 9px;
  letter-spacing: 0.2em;
  text-transform: uppercase;
  color: var(--ochre);
  opacity: 0.55;
  margin: 20px 0 8px;
}}
.toc-list {{ list-style: none; display: flex; flex-direction: column; gap: 1px; }}
.toc-list a {{
  display: block;
  font-family: 'Cormorant Garamond', serif;
  font-size: 13.5px;
  line-height: 1.45;
  color: var(--ink);
  opacity: 0.65;
  text-decoration: none;
  padding: 4px 8px 4px 10px;
  border-left: 2px solid transparent;
  transition: all 0.15s ease;
}}
.toc-list a:hover, .toc-list a.active {{
  opacity: 1;
  border-left-color: var(--ochre);
  color: var(--sienna);
}}
.toc-list a .toc-num {{
  font-family: 'Source Code Pro', monospace;
  font-size: 11px;
  color: var(--ochre);
  opacity: 0.7;
  margin-right: 5px;
}}
.sidebar-rule {{ height: 1px; background: var(--rule); margin: 20px 0; }}
.sidebar-meta {{
  font-family: 'Source Code Pro', monospace;
  font-size: 8.5px;
  letter-spacing: 0.06em;
  color: var(--ink);
  opacity: 0.4;
  line-height: 1.8;
  margin-top: auto;
  padding-top: 20px;
}}
.sidebar-meta a {{ color: var(--deep-ochre); text-decoration: none; }}

/* ============================================================
   MAIN
   ============================================================ */
.main {{ flex: 1; min-width: 0; padding: 0 0 80px; }}

/* ============================================================
   COVER
   ============================================================ */
.cover {{
  background: linear-gradient(170deg, #1a1a0e 0%, #0f1d35 60%, #1e2d4a 100%);
  padding: 72px 64px 64px;
  position: relative;
  overflow: hidden;
}}
.cover::before {{
  content: '';
  position: absolute; inset: 0;
  background: radial-gradient(ellipse at 30% 50%, rgba(200,131,42,0.08) 0%, transparent 65%);
}}
.cover::after {{
  content: '';
  position: absolute;
  bottom: 0; left: 64px; right: 64px;
  height: 1px;
  background: linear-gradient(to right, transparent, var(--gold), transparent);
  opacity: 0.3;
}}
.cover-inner {{ position: relative; z-index: 1; max-width: var(--max-body); }}
.cover-logo {{
  height: 64px; width: auto;
  margin-bottom: 24px;
  opacity: 0.92;
  display: block;
}}
.cover-org {{
  font-family: 'Source Code Pro', monospace;
  font-size: 12.5px;
  letter-spacing: 0.25em;
  text-transform: uppercase;
  color: var(--gold);
  opacity: 0.6;
  margin-bottom: 28px;
}}
.cover-title {{
  font-family: 'Cinzel', serif;
  font-size: clamp(22px, 3.5vw, 34px);
  font-weight: 500;
  letter-spacing: 0.04em;
  color: #f5f0e6;
  line-height: 1.25;
  margin-bottom: 16px;
}}
.cover-subtitle {{
  font-family: 'Cormorant Garamond', serif;
  font-style: italic;
  font-size: 19px;
  color: var(--pale-ochre);
  opacity: 0.75;
  line-height: 1.5;
  margin-bottom: 40px;
  max-width: 600px;
}}
.cover-meta {{
  display: flex;
  gap: 38px;
  flex-wrap: wrap;
  border-top: 1px solid rgba(212,168,67,0.15);
  padding-top: 24px;
}}
.meta-label {{
  font-family: 'Source Code Pro', monospace;
  font-size: 12px;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  color: var(--gold);
  opacity: 0.5;
  margin-bottom: 3px;
}}
.meta-value {{
  font-family: 'Cormorant Garamond', serif;
  font-size: 15px;
  color: var(--pale-ochre);
  opacity: 0.8;
}}

/* ============================================================
   DOCUMENT BODY
   ============================================================ */
.doc-body {{ max-width: calc(var(--max-body) + 128px); padding: 0 64px; }}

/* Section */
.doc-section {{
  padding: 56px 0 0;
  border-top: 1px solid var(--rule);
  margin-top: 0;
  animation: fadeIn 0.4s ease both;
}}
.doc-section:first-of-type {{ border-top: none; padding-top: 52px; }}
@keyframes fadeIn {{
  from {{ opacity: 0; transform: translateY(8px); }}
  to   {{ opacity: 1; transform: translateY(0); }}
}}
.section-eyebrow {{
  font-family: 'Source Code Pro', monospace;
  font-size: 8px;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  color: var(--ochre);
  opacity: 0.65;
  margin-bottom: 10px;
}}
h2.section-title {{
  font-family: 'Cinzel', serif;
  font-size: 24px;
  font-weight: 500;
  letter-spacing: 0.05em;
  color: var(--indigo);
  line-height: 1.3;
  margin-bottom: 28px;
}}
h3.subsection-title {{
  font-family: 'Cinzel', serif;
  font-size: 15px;
  font-weight: 500;
  letter-spacing: 0.1em;
  color: var(--sienna);
  margin: 32px 0 14px;
  padding-bottom: 6px;
  border-bottom: 1px solid var(--rule);
}}
h4.subsubsection-title {{
  font-family: 'Cormorant Garamond', serif;
  font-size: 17px;
  font-weight: 600;
  font-style: italic;
  color: var(--deep-ochre);
  margin: 24px 0 10px;
}}

/* ============================================================
   PROSE
   ============================================================ */
.doc-body p {{
  font-size: 19px;
  line-height: 1.8;
  color: var(--ink);
  margin-bottom: 20px;
  max-width: var(--max-body);
}}
.doc-body p + p {{ text-indent: 1.5em; }}
.doc-body p:first-of-type {{ text-indent: 0; }}
.doc-body strong {{ font-weight: 600; }}
.doc-body em {{ font-style: italic; }}
.doc-body code {{
  font-family: 'Source Code Pro', monospace;
  font-size: 14px;
  background: rgba(200,131,42,0.08);
  color: var(--deep-ochre);
  padding: 1px 5px;
  border-radius: 2px;
}}

/* ============================================================
   LISTS
   ============================================================ */
.doc-list {{
  margin: 4px 0 20px 0;
  padding-left: 28px;
  max-width: var(--max-body);
}}
.doc-list li {{
  font-size: 18px;
  line-height: 1.75;
  margin-bottom: 6px;
  color: var(--ink);
}}
.doc-list-num {{ list-style: decimal; }}
ul.doc-list {{ list-style: none; }}
ul.doc-list li::before {{
  content: "—";
  color: var(--ochre);
  opacity: 0.6;
  margin-right: 10px;
  margin-left: -24px;
  font-weight: 300;
}}

/* ============================================================
   CODE BLOCK
   ============================================================ */
.code-block {{
  background: #0f1a2e;
  border: 1px solid rgba(212,168,67,0.15);
  border-radius: 3px;
  padding: 24px 26px;
  margin: 24px 0;
  max-width: var(--max-body);
  overflow-x: auto;
}}
.code-block pre {{
  font-family: 'Source Code Pro', monospace;
  font-size: 11.5px;
  line-height: 1.7;
  color: #c8d8e8;
  white-space: pre;
  overflow-x: auto;
  margin: 0;
}}

/* ============================================================
   NOTE / WARNING / TIP
   ============================================================ */
.doc-note, .doc-warning, .doc-tip {{
  margin: 20px 0;
  padding: 14px 20px;
  border-radius: 2px;
  max-width: var(--max-body);
  font-size: 16px;
  line-height: 1.7;
}}
.doc-note {{
  background: rgba(30,45,74,0.05);
  border-left: 3px solid var(--indigo);
}}
.doc-warning {{
  background: rgba(192,0,0,0.04);
  border-left: 3px solid #c00000;
}}
.doc-tip {{
  background: rgba(45,74,45,0.05);
  border-left: 3px solid var(--forest);
}}
.note-label {{
  font-family: 'Source Code Pro', monospace;
  font-size: 9px;
  letter-spacing: 0.18em;
  text-transform: uppercase;
  font-weight: 600;
  margin-right: 8px;
}}
.doc-note .note-label    {{ color: var(--indigo); }}
.doc-warning .note-label {{ color: #c00000; }}
.doc-tip .note-label     {{ color: var(--forest); }}

/* ============================================================
   TABLE
   ============================================================ */
.module-table-wrapper {{ margin: 28px 0; overflow-x: auto; max-width: var(--max-body); }}
table.module-table {{ width: 100%; border-collapse: collapse; font-size: 12.5px; }}
.module-table th {{
  font-family: 'Source Code Pro', monospace;
  font-size: 7.5px;
  letter-spacing: 0.15em;
  text-transform: uppercase;
  color: var(--ochre);
  opacity: 0.7;
  text-align: left;
  padding: 8px 12px;
  border-bottom: 1px solid var(--rule-strong);
  white-space: nowrap;
}}
.module-table td {{
  padding: 8px 12px;
  border-bottom: 1px solid var(--rule);
  vertical-align: top;
  font-size: 13px;
  line-height: 1.55;
}}
.module-table tr:hover td {{ background: rgba(200,131,42,0.03); }}

/* ============================================================
   HORIZONTAL RULE
   ============================================================ */
.doc-rule {{
  border: none;
  height: 1px;
  background: var(--rule);
  margin: 40px 0;
  max-width: var(--max-body);
}}

/* ============================================================
   FOOTER
   ============================================================ */
.doc-footer {{
  margin-top: 60px;
  padding: 28px 0;
  border-top: 1px solid var(--rule);
  max-width: var(--max-body);
  display: flex;
  justify-content: space-between;
  align-items: flex-end;
  gap: 24px;
  flex-wrap: wrap;
}}
.footer-org {{
  font-family: 'Cinzel', serif;
  font-size: 9px;
  font-weight: 400;
  letter-spacing: 0.22em;
  text-transform: uppercase;
  color: var(--ochre);
  opacity: 0.6;
}}
.footer-meta {{
  font-family: 'Source Code Pro', monospace;
  font-size: 8px;
  letter-spacing: 0.06em;
  color: var(--ink);
  opacity: 0.35;
}}

/* ============================================================
   MOBILE NAV
   ============================================================ */
.mobile-nav {{
  display: none;
  position: sticky;
  top: 0; z-index: 100;
  background: #f0ebe0;
  border-bottom: 1px solid var(--rule);
  padding: 12px 20px;
  gap: 12px;
  align-items: center;
}}
.mobile-nav-title {{
  font-family: 'Cinzel', serif;
  font-size: 10px;
  letter-spacing: 0.1em;
  color: var(--indigo);
  flex: 1;
}}
.mobile-toc-toggle {{
  font-family: 'Source Code Pro', monospace;
  font-size: 8px;
  letter-spacing: 0.12em;
  text-transform: uppercase;
  color: var(--ochre);
  background: none;
  border: 1px solid var(--rule-strong);
  padding: 5px 12px;
  cursor: pointer;
  border-radius: 2px;
}}
.mobile-toc-panel {{
  display: none; position: fixed; inset: 0; z-index: 200;
  background: rgba(26,18,8,0.7);
}}
.mobile-toc-panel.open {{ display: block; }}
.mobile-toc-inner {{
  position: absolute; top: 0; left: 0;
  width: 280px; height: 100%;
  background: #f0ebe0;
  padding: 28px 20px;
  overflow-y: auto;
}}
.mobile-toc-close {{
  position: absolute; top: 16px; right: 16px;
  background: none; border: none;
  font-size: 18px; cursor: pointer;
  color: var(--ink); opacity: 0.5;
}}

/* ============================================================
   RESPONSIVE
   ============================================================ */
@media (max-width: 900px) {{
  .sidebar {{ display: none; }}
  .mobile-nav {{ display: flex; }}
  .doc-body {{ padding: 0 28px; }}
  .cover {{ padding: 48px 28px; }}
}}
@media (max-width: 600px) {{
  .doc-body {{ padding: 0 18px; }}
  .cover {{ padding: 36px 18px; }}
}}

/* ============================================================
   PRINT
   ============================================================ */
@media print {{
  .sidebar, .mobile-nav {{ display: none !important; }}
  body::before {{ display: none; }}
  .main {{ padding: 0; }}
  .cover {{ print-color-adjust: exact; -webkit-print-color-adjust: exact; }}
  .doc-body {{ padding: 0 48px; }}
  .doc-section {{ page-break-inside: avoid; }}
  h2.section-title {{ page-break-after: avoid; }}
  .code-block {{ page-break-inside: avoid; }}
  .layout {{ display: block; }}
}}
</style>
</head>
<body>

<!-- MOBILE NAV -->
<div class="mobile-nav" id="mobileNav">
  <div class="mobile-nav-title">{escape(self.title)}</div>
  <button class="mobile-toc-toggle" onclick="toggleMobileToc()">Contents</button>
</div>
<div class="mobile-toc-panel" id="mobileTocPanel">
  <div class="mobile-toc-inner">
    <button class="mobile-toc-close" onclick="toggleMobileToc()">&#x2715;</button>
    <div style="font-family:'Cinzel',serif;font-size:10px;letter-spacing:0.12em;color:var(--ochre);margin-bottom:20px;">Contents</div>
    <ul class="toc-list" id="mobileTocList"></ul>
  </div>
</div>

<div class="layout">

  <!-- SIDEBAR -->
  <aside class="sidebar" id="sidebar">
    <img src="assets/IHS-Logo.jpg" alt="Iroko Historical Society" class="sidebar-logo-img"
         onerror="this.style.display='none'">
    <div class="sidebar-org">Iroko Historical Society</div>
    <div class="sidebar-title">{escape(self.title)}</div>
    <div class="sidebar-version">{escape(self.doc_type)}{ver_display}</div>

    <div class="toc-section-head">Contents</div>
    <ul class="toc-list" id="tocList">
{toc_html}
    </ul>

    <div class="sidebar-rule"></div>

    <div class="sidebar-meta">
      irokosociety.org<br>
      ontology.irokosociety.org
    </div>
  </aside>

  <!-- MAIN -->
  <main class="main">

    <!-- COVER -->
    <div class="cover">
      <div class="cover-inner">
        <img src="assets/IHS-Logo.jpg" alt="Iroko Historical Society" class="cover-logo"
             onerror="this.style.display='none'">
        <div class="cover-org">Iroko Historical Society &middot; {escape(self.doc_type)}</div>
        <h1 class="cover-title">{escape(self.title)}</h1>
        {subtitle_html}
        <div class="cover-meta">
          <div class="cover-meta-item">
            <div class="meta-label">Publisher</div>
            <div class="meta-value">{escape(self.org)}</div>
          </div>
          <div class="cover-meta-item">
            <div class="meta-label">Author</div>
            <div class="meta-value">{escape(self.author)}</div>
          </div>
          {"" if not self.version else f'<div class="cover-meta-item"><div class="meta-label">Version</div><div class="meta-value">{escape(self.version)}</div></div>'}
          <div class="cover-meta-item">
            <div class="meta-label">Site</div>
            <div class="meta-value">
              <a href="https://irokosociety.org" style="color:var(--pale-ochre);text-decoration:none;">irokosociety.org</a>
            </div>
          </div>
        </div>
      </div>
    </div>

    <!-- BODY -->
    <div class="doc-body">
{body_html}

      <!-- FOOTER -->
      <footer class="doc-footer">
        <div class="footer-org">Iroko Historical Society</div>
        <div class="footer-meta">irokosociety.org &middot; ontology.irokosociety.org</div>
      </footer>
    </div>

  </main>
</div>

<script>
// ============================================================
// TOC ACTIVE HIGHLIGHTING
// ============================================================
(function() {{
  const sections = document.querySelectorAll('.doc-section, h3.subsection-title');
  const tocLinks = document.querySelectorAll('#tocList a');
  const mobileTocLinks = document.querySelectorAll('#mobileTocList a');

  // Clone desktop TOC into mobile TOC
  const mobileTocList = document.getElementById('mobileTocList');
  const desktopTocList = document.getElementById('tocList');
  if (desktopTocList && mobileTocList) {{
    mobileTocList.innerHTML = desktopTocList.innerHTML;
  }}

  const allTocLinks = document.querySelectorAll('.toc-list a');

  const observer = new IntersectionObserver((entries) => {{
    entries.forEach(entry => {{
      if (entry.isIntersecting) {{
        const id = entry.target.id;
        allTocLinks.forEach(a => {{
          a.classList.toggle('active', a.getAttribute('href') === '#' + id);
        }});
      }}
    }});
  }}, {{ rootMargin: '-15% 0px -75% 0px' }});

  document.querySelectorAll('[id]').forEach(el => observer.observe(el));
}})();

// ============================================================
// MOBILE TOC TOGGLE
// ============================================================
function toggleMobileToc() {{
  const panel = document.getElementById('mobileTocPanel');
  panel.classList.toggle('open');
}}
document.getElementById('mobileTocPanel').addEventListener('click', function(e) {{
  if (e.target === this) toggleMobileToc();
}});
document.querySelectorAll('#mobileTocList a').forEach(a => {{
  a.addEventListener('click', toggleMobileToc);
}});
</script>
</body>
</html>"""


# ============================================================
# CLI
# ============================================================

def main():
    parser = argparse.ArgumentParser(
        description="Convert a .docx file to an Iroko-styled HTML document."
    )
    parser.add_argument("input", help="Path to the input .docx file")
    parser.add_argument("output", nargs="?", help="Path for the output .html file (optional)")
    parser.add_argument("--title",    help="Override document title")
    parser.add_argument("--subtitle", help="Cover subtitle / tagline")
    parser.add_argument("--version",  help="Version string (e.g. v1.0, 2026)")
    parser.add_argument("--author",   default="Iroko Historical Society", help="Author name")
    parser.add_argument("--org",      default="Iroko Historical Society", help="Publishing organization")
    parser.add_argument("--type",     default="Technical Manual",
                        dest="doc_type", help="Document type label for cover (default: Technical Manual)")
    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"ERROR: File not found: {args.input}")
        sys.exit(1)

    output = args.output or (Path(args.input).stem + ".html")

    print(f"Converting: {args.input}")
    converter = DocxConverter(
        args.input,
        title    = args.title,
        subtitle = args.subtitle,
        version  = args.version,
        author   = args.author,
        org      = args.org,
        doc_type = args.doc_type,
    )
    html = converter.render()

    with open(output, "w", encoding="utf-8") as f:
        f.write(html)

    headings_found = len(converter.headings)
    print(f"Done. Output: {output}")
    print(f"  Headings indexed: {headings_found}")
    print(f"  Title:  {converter.title}")
    if converter.subtitle:
        print(f"  Subtitle: {converter.subtitle}")


if __name__ == "__main__":
    main()
